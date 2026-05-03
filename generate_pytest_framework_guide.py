"""Generate ReportStack_pytest_Framework_Guide.docx.

A pytest-native framework guide for testing Linux radio devices with
ReportStack reporting. Companion to test_framework_starter/.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

BRAND_DEEP = RGBColor(0x1E, 0x40, 0xAF)
BRAND_DARK = RGBColor(0x0F, 0x17, 0x2A)
BRAND_MUTED = RGBColor(0x64, 0x74, 0x8B)
BRAND_LIGHT = RGBColor(0x94, 0xA3, 0xB8)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = BRAND_DARK if level == 1 else BRAND_DEEP
    hs.font.bold = True


def add_table(headers, rows, widths_in=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(10)
    for row in rows:
        r = t.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = str(val)
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if widths_in:
        for row in t.rows:
            for i, w in enumerate(widths_in):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def add_callout(text, kind="tip"):
    label_map = {
        "tip": ("Tip", "DBEAFE"),
        "warn": ("Watch out", "FEF3C7"),
        "note": ("Note", "F1F5F9"),
    }
    label, fill = label_map.get(kind, label_map["note"])
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill,
    })
    tcPr.append(shd)
    cell.text = ""
    p = cell.paragraphs[0]
    lr = p.add_run(f"{label}  "); lr.bold = True; lr.font.color.rgb = BRAND_DEEP
    br = p.add_run(text); br.font.size = Pt(10.5)
    cell.width = Inches(6.5)
    doc.add_paragraph()


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)


def add_code(text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F1F5F9",
    })
    tcPr.append(shd)
    cell.text = ""
    for i, line in enumerate(text.splitlines()):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = BRAND_DARK
    cell.width = Inches(6.5)
    doc.add_paragraph()


def section_break():
    doc.add_page_break()


# ============================================================
# COVER
# ============================================================
for _ in range(3):
    doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ReportStack")
r.font.size = Pt(40); r.bold = True; r.font.color.rgb = BRAND_DARK
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("pytest framework guide")
r.font.size = Pt(22); r.font.color.rgb = BRAND_DEEP
doc.add_paragraph()
desc = doc.add_paragraph(); desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = desc.add_run("A pytest-native framework for testing Linux radio devices with ReportStack reporting")
r.font.size = Pt(13); r.font.color.rgb = BRAND_MUTED
doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    f"Version 1.0\n{datetime.date.today().strftime('%B %d, %Y')}\n\n"
    "Companion to test_framework_starter/. Replaces the Java/TestNG\n"
    "Osprey framework analysis with a pytest-native architecture."
)
r.font.size = Pt(11); r.font.color.rgb = BRAND_LIGHT
section_break()


# ============================================================
# CONTENTS
# ============================================================
doc.add_heading("Contents", level=1)
toc = [
    "1. Overview",
    "2. Tech stack",
    "3. Project structure",
    "4. Architecture and design patterns",
    "5. Dependency management",
    "6. Pytest fundamentals (markers, fixtures, conftest, parametrize, xdist)",
    "7. Test categories and marker taxonomy",
    "8. Linux radio device control (SSH and serial)",
    "9. ReportStack integration",
    "10. AI failure analysis",
    "11. Defects and external links",
    "12. Performance KPIs",
    "13. Artifacts (MinIO via ReportStack)",
    "14. Notifications",
    "15. Jenkins CI/CD pipeline",
    "16. Migration mapping: Osprey → ReportStack pytest",
    "17. Worked example (one full test, end to end)",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
section_break()


# ============================================================
# 1. OVERVIEW
# ============================================================
doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "This guide describes a pytest-native automation framework for Linux radio "
    "devices, with results reported through ReportStack. It draws inspiration "
    "from the Osprey/TestNG framework analysis (where the layered architecture "
    "and team-based suite organization come from) but replaces every Java, "
    "Maven, TestNG, and AspectJ component with its idiomatic pytest equivalent."
)
doc.add_paragraph("The framework deliberately keeps four invariants from Osprey:")
add_bullets([
    "Separation of concerns — Tests → Steps → Devices → Commands.",
    "Configuration over code — JSON command catalogs, marker-based suite selection, runtime CLI flags for device endpoints.",
    "Team-based test organization — markers like team_allstars and team_gogeta let multiple teams share one repo.",
    "Multi-level reporting — execution logs + screenshots + diagnostic snapshots + AI classification, all visible in the ReportStack dashboard.",
])
doc.add_paragraph("And changes everything else:")
add_bullets([
    "Java 11 → Python 3.12.",
    "TestNG XML suites → pytest markers + pytest.ini.",
    "Maven + AspectJ → pyproject.toml + pytest hooks.",
    "Appium / UIAutomator2 / ADB → SSH (paramiko) and serial (pyserial) for radios.",
    "ReportPortal listener → the pytest-automation-reports plugin shipped in this repo.",
    "AWS Lambda triage → Ollama (local LLM) via ReportStack's analyzer.",
])

section_break()

# ============================================================
# 2. TECH STACK
# ============================================================
doc.add_heading("2. Tech stack", level=1)
add_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Language", "Python 3.12", "Backend match; type hints throughout"],
        ["Test runner", "pytest ≥ 7.4", "Markers, fixtures, parametrize, plugins"],
        ["Parallelism", "pytest-xdist ≥ 3.5", "-n N replaces TestNG parallel=methods"],
        ["Retry", "pytest-rerunfailures ≥ 14", "Equivalent of TestNG RetryListener"],
        ["Build / packaging", "pyproject.toml + setuptools", "Replaces Maven pom.xml"],
        ["Device — SSH", "paramiko ≥ 3.4", "Linux radio control plane over SSH"],
        ["Device — serial", "pyserial ≥ 3.5", "Console-only / bring-up scenarios"],
        ["Logging", "structlog ≥ 24", "Structured logs forwarded to ReportStack"],
        ["Reporting", "pytest-automation-reports", "In-repo plugin; replaces RP listener"],
        ["AI triage", "Ollama (local LLM)", "Replaces AWS Lambda; local, no API keys"],
        ["Issue tracking", "ReportStack defects + external_link", "Replaces Jira/QMetry TAM API"],
        ["Artifact storage", "ReportStack attachments → MinIO", "Replaces Nexus binary repo"],
        ["Image / screenshots", "Pillow ≥ 10", "Diff and processing"],
        ["CI/CD", "Jenkins (existing)", "Pipeline calls pytest instead of mvn"],
    ],
    widths_in=[1.4, 2.0, 3.1],
)
add_callout(
    "Pytest doesn't need AOP. The cross-cutting concerns Osprey solved with "
    "AspectJ (logging, retry, reporting) are first-class pytest features: "
    "hooks for cross-cutting, fixtures for setup, plugins for extension.",
    "tip",
)
section_break()

# ============================================================
# 3. PROJECT STRUCTURE
# ============================================================
doc.add_heading("3. Project structure", level=1)
doc.add_paragraph(
    "The starter project layout — see test_framework_starter/ in this repo "
    "for the actual files."
)
add_code(
    "test_framework_starter/\n"
    "├── pyproject.toml          # deps + project metadata\n"
    "├── pytest.ini              # marker registry + addopts\n"
    "├── conftest.py             # root fixtures + ReportStack hooks\n"
    "├── radios/                 # device abstractions\n"
    "│   ├── base.py             # RadioDevice abstract base + CommandResult\n"
    "│   ├── ssh_radio.py        # paramiko-backed SSHRadio\n"
    "│   └── serial_radio.py     # pyserial-backed SerialRadio\n"
    "├── commands/               # JSON command catalogs\n"
    "│   ├── loader.py           # validates required keys\n"
    "│   └── examples/\n"
    "│       ├── wnc_radio.json\n"
    "│       └── generic_linux.json\n"
    "├── steps/                  # composable business operations\n"
    "│   ├── power.py            # bring_up, shut_down\n"
    "│   ├── tune.py             # tune_and_wait_for_lock\n"
    "│   └── measure.py          # measure_rssi_average\n"
    "└── tests/\n"
    "    ├── conftest.py         # test-scoped fixtures (radio_up)\n"
    "    ├── test_smoke.py       # smoke + parametrized tune tests\n"
    "    ├── test_regression.py  # regression tests with team markers\n"
    "    └── test_kpi.py         # lock-time KPI w/ threshold assertion"
)
add_callout(
    "tests/ is a flat folder. Per-team or per-feature breakdowns happen via "
    "markers (team_allstars, feature_tune), not directories. This avoids the "
    "duplication-by-folder problem Osprey ran into with 70+ feature suite XMLs.",
    "tip",
)
section_break()

# ============================================================
# 4. ARCHITECTURE & PATTERNS
# ============================================================
doc.add_heading("4. Architecture and design patterns", level=1)
doc.add_paragraph(
    "Four layers, top to bottom. Every layer depends only on the one below — "
    "the test layer never reaches into a transport client; the radio layer "
    "never knows about pytest."
)
add_table(
    ["Layer", "Responsibility", "Pattern"],
    [
        ["Tests (pytest functions)", "Assertions and data-driven scenarios", "Specification by example"],
        ["Steps (Python modules)", "Business-flavoured verbs (bring_up, tune_and_wait_for_lock)", "Strategy + Builder"],
        ["Radios (RadioDevice subclasses)", "Device control — power, tune, read_rssi", "Page Object Model"],
        ["Commands (JSON catalogs)", "Per-device-family text strings sent over the wire", "Factory + lookup"],
    ],
    widths_in=[2.2, 2.8, 1.5],
)
doc.add_heading("Design patterns", level=2)
add_table(
    ["Pattern", "Where it lives", "Why"],
    [
        ["Page Object Model", "radios/base.py + impls", "Tests call device.tune(), not a raw shell command"],
        ["Factory + lookup", "commands/loader.py", "JSON files swap in per-device-family strings without code change"],
        ["Strategy", "SSHRadio vs SerialRadio", "Same RadioDevice surface, different transport"],
        ["Template Method", "RadioDevice.session() context manager", "Standard cleanup path; subclasses fill in run/close"],
        ["Listener / Observer", "pytest hooks + the AR plugin", "Cross-cutting — logging, capture-on-fail, reporting"],
        ["Retry", "pytest-rerunfailures + retry_of in ReportStack", "Surface flakes without losing them"],
    ],
    widths_in=[1.7, 2.4, 2.4],
)
section_break()

# ============================================================
# 5. DEPENDENCY MANAGEMENT
# ============================================================
doc.add_heading("5. Dependency management", level=1)
doc.add_paragraph(
    "pyproject.toml replaces pom.xml. Optional groups (test, dev) replace Maven "
    "profiles. Pin specific versions for reproducibility, looser ranges for libs."
)
add_code(
    "[project]\n"
    "name = \"radio-test-framework-starter\"\n"
    "requires-python = \">=3.10\"\n"
    "dependencies = [\n"
    "  \"pytest>=7.4\",\n"
    "  \"pytest-xdist>=3.5\",\n"
    "  \"pytest-rerunfailures>=14.0\",\n"
    "  \"paramiko>=3.4\",\n"
    "  \"pyserial>=3.5\",\n"
    "  \"Pillow>=10.0\",\n"
    "]\n"
    "\n"
    "[project.optional-dependencies]\n"
    "dev = [\"ruff>=0.5\", \"mypy>=1.10\"]"
)
add_callout(
    "Install in editable mode during development (pip install -e .) so changes "
    "to radios/, steps/, and commands/ are picked up without reinstall.",
    "note",
)
section_break()

# ============================================================
# 6. PYTEST FUNDAMENTALS
# ============================================================
doc.add_heading("6. Pytest fundamentals", level=1)
doc.add_paragraph(
    "If you're coming from TestNG, the cheat sheet:"
)
add_table(
    ["TestNG concept", "Pytest equivalent"],
    [
        ["@Test(groups = {\"smoke\"})", "@pytest.mark.smoke (registered in pytest.ini)"],
        ["@DataProvider", "@pytest.mark.parametrize"],
        ["@BeforeSuite / @AfterSuite", "Session-scoped fixture (yield-style)"],
        ["@BeforeMethod / @AfterMethod", "Function-scoped fixture (yield-style)"],
        ["BaseTest inheritance", "conftest.py + composable fixtures (no inheritance)"],
        ["RetryListener", "pytest-rerunfailures (--reruns N)"],
        ["FailListener", "pytest_runtest_makereport hook in conftest"],
        ["TestNG XML suite", "pytest.ini markers + pytest -m expression"],
        ["parallel=methods, thread-count=2", "pytest-xdist: pytest -n 2"],
        ["@Listeners", "pytest plugin (entry point or conftest.py)"],
    ],
    widths_in=[2.7, 3.8],
)

doc.add_heading("Markers", level=2)
doc.add_paragraph(
    "Markers tag tests with metadata. Register every marker in pytest.ini — "
    "the strict-markers option turns typos into collection errors instead of "
    "silently-skipped tests."
)
add_code(
    "# pytest.ini\n"
    "[pytest]\n"
    "markers =\n"
    "    smoke: fast, broad coverage; runs on every push\n"
    "    regression: full functional suite; runs nightly\n"
    "    kpi: performance KPI measurements\n"
    "    team_allstars: AllStars team's regression scope\n"
    "    feature_tune: tuning / channel-change feature\n"
    "addopts =\n"
    "    --strict-markers\n"
    "    -ra"
)

doc.add_heading("Fixtures (the Page Object enabler)", level=2)
doc.add_paragraph(
    "Fixtures replace TestNG's @BeforeMethod/@AfterMethod and the implicit "
    "field state of a BaseTest subclass. Yield-style fixtures handle setup "
    "and teardown in one block."
)
add_code(
    "import pytest\n"
    "from radios.ssh_radio import SSHRadio\n"
    "from commands import load_catalog\n"
    "\n"
    "@pytest.fixture(scope=\"session\")\n"
    "def radio(request):\n"
    "    catalog = load_catalog(request.config.getoption(\"--device-cmds\"))\n"
    "    device = SSHRadio(\n"
    "        name=request.config.getoption(\"--device-host\"),\n"
    "        host=request.config.getoption(\"--device-host\"),\n"
    "        username=request.config.getoption(\"--device-user\"),\n"
    "        command_catalog=catalog,\n"
    "    )\n"
    "    with device.session() as d:\n"
    "        yield d   # teardown after this line"
)

doc.add_heading("Parametrize (replaces @DataProvider)", level=2)
add_code(
    "@pytest.mark.smoke\n"
    "@pytest.mark.parametrize(\n"
    "    \"freq_mhz\",\n"
    "    [88.5, 100.0, 433.92, 868.0],\n"
    "    ids=[\"fm_low\", \"fm_mid\", \"ism_433\", \"ism_868\"],\n"
    ")\n"
    "def test_tune_locks_within_5s(radio_up, freq_mhz):\n"
    "    elapsed = tune_and_wait_for_lock(radio_up, freq_mhz, timeout_seconds=5.0)\n"
    "    assert elapsed < 5.0"
)

doc.add_heading("Parallel runs with pytest-xdist", level=2)
add_code(
    "# Two workers; one radio per worker via separate --device-host values.\n"
    "pytest -n 2 -m smoke \\\n"
    "  --device-host=radio1.lab \\\n"
    "  --device-host=radio2.lab"
)
add_callout(
    "xdist runs tests across multiple workers. For radio testing each worker "
    "needs its own physical device — pass the same flag twice, or read a "
    "device pool from a config file in conftest.",
    "warn",
)

section_break()

# ============================================================
# 7. MARKERS
# ============================================================
doc.add_heading("7. Test categories and marker taxonomy", level=1)
doc.add_paragraph(
    "Three orthogonal axes — kind, team, feature. Combine them with boolean "
    "expressions on the command line."
)
add_table(
    ["Axis", "Markers", "Example"],
    [
        ["Kind", "smoke, regression, kpi, nfr, stress", "pytest -m smoke"],
        ["Team", "team_allstars, team_gogeta, team_xteam", "pytest -m \"regression and team_allstars\""],
        ["Feature", "feature_tune, feature_power, feature_measure", "pytest -m \"smoke and feature_tune\""],
        ["Speed", "slow", "pytest -m \"regression and not slow\""],
    ],
    widths_in=[1.0, 2.7, 2.8],
)
add_callout(
    "Stop creating new test files for each suite. With markers, one regression.py "
    "can serve every team via @pytest.mark.team_xxx.",
    "tip",
)

section_break()

# ============================================================
# 8. RADIO CONTROL
# ============================================================
doc.add_heading("8. Linux radio device control", level=1)
doc.add_paragraph(
    "Radios speak shell or serial. The RadioDevice abstract base normalises "
    "both into a single API; tests don't know which transport they're using."
)
doc.add_heading("The RadioDevice contract", level=2)
add_code(
    "class RadioDevice(ABC):\n"
    "    @abstractmethod\n"
    "    def run(self, command: str, *, timeout: float = 10.0) -> CommandResult: ...\n"
    "    @abstractmethod\n"
    "    def close(self) -> None: ...\n"
    "\n"
    "    # Catalog-driven helpers — same on every transport\n"
    "    def cmd(self, key: str, **fmt) -> CommandResult: ...\n"
    "    def power_on(self) -> None: ...\n"
    "    def tune(self, freq_mhz: float) -> None: ...\n"
    "    def read_rssi(self) -> int: ...\n"
    "    def is_locked(self) -> bool: ..."
)

doc.add_heading("JSON command catalogs", level=2)
doc.add_paragraph(
    "Each device family ships with a JSON file mapping logical command names "
    "to the actual shell strings. Adding a new device is a matter of writing "
    "one JSON file — no Python code changes."
)
add_code(
    "// commands/examples/wnc_radio.json\n"
    "{\n"
    "  \"power_on\":        \"radioctl --on\",\n"
    "  \"power_off\":       \"radioctl --off\",\n"
    "  \"set_freq\":        \"radioctl --freq {mhz}\",\n"
    "  \"read_rssi\":       \"radioctl --query rssi\",\n"
    "  \"read_lock_state\": \"radioctl --query lock\"\n"
    "}"
)
add_callout(
    "loader.py validates that every catalog defines the five required keys "
    "(power_on, power_off, set_freq, read_rssi, read_lock_state). Missing a "
    "key is a collection-time error, not a runtime surprise.",
    "note",
)

doc.add_heading("Picking a transport", level=2)
add_table(
    ["Use", "When"],
    [
        ["SSHRadio", "Device has Ethernet/WiFi and runs sshd"],
        ["SerialRadio", "Bring-up / lab bench, or no network stack"],
        ["Future: TCP socket / SCPI", "Test instruments speaking SCPI over LAN — implement RadioDevice the same way"],
    ],
    widths_in=[2.0, 4.5],
)

section_break()

# ============================================================
# 9. REPORTSTACK INTEGRATION
# ============================================================
doc.add_heading("9. ReportStack integration", level=1)
doc.add_paragraph(
    "The pytest-automation-reports plugin (in plugins/ at the repo root) "
    "registers as a pytest plugin and forwards every test outcome to "
    "ReportStack via REST. Install it once, then enable per run with CLI flags."
)
add_code("pip install -e plugins/pytest-automation-reports/")
doc.add_heading("CLI flags", level=2)
add_table(
    ["Flag", "Default", "Purpose"],
    [
        ["--ar-url", "(none — required to enable)", "Backend base URL, e.g. http://reports.local:8000/api/v1"],
        ["--ar-launch-name", "auto-generated timestamp", "Label shown on the launches list"],
        ["--ar-launch-description", "(none)", "Free-form context — build number, branch, etc."],
        ["--ar-tag", "(none)", "Repeatable tag for filtering (e.g. nightly, hotfix)"],
        ["--ar-auto-analyze", "false", "Trigger AI analysis as soon as the launch finishes"],
    ],
    widths_in=[2.0, 2.0, 2.5],
)

doc.add_heading("What the plugin does for you", level=2)
add_bullets([
    "Creates a Launch on session start; finishes it on session end with PASSED / FAILED / STOPPED.",
    "Reports each test as a TestItem with status, duration, error message, and stack trace.",
    "Captures pytest log records and uploads them as TestLogs in batches.",
    "Uploads files registered via request.node.user_properties[(\"attachment\", path)] as Attachments.",
    "Forwards numeric KPIs registered as user_properties[(\"kpi:name\", value)] for the Trends page.",
])

section_break()

# ============================================================
# 10. AI ANALYSIS
# ============================================================
doc.add_heading("10. AI failure analysis", level=1)
doc.add_paragraph(
    "When --ar-auto-analyze is on, ReportStack triggers Ollama-backed "
    "classification on every failed test once the launch finishes. The "
    "analyzer reads the test's error message, last 50 stack-trace lines, and "
    "last 20 log entries; returns one of PRODUCT_BUG, AUTOMATION_BUG, "
    "SYSTEM_ISSUE, NO_DEFECT, TO_INVESTIGATE."
)
add_callout(
    "If Ollama is unreachable, the analyzer falls back to TO_INVESTIGATE "
    "rather than crashing the BackgroundTask — a safety net added in May 2026.",
    "tip",
)
doc.add_heading("Manual override", level=2)
doc.add_paragraph(
    "On the Test Detail page, the Defect Selector dropdown lets a triager "
    "override the AI's call. The override is stored as a separate "
    "FailureAnalysis row with source=MANUAL, so the original AI decision is "
    "preserved for audit."
)

section_break()

# ============================================================
# 11. DEFECTS
# ============================================================
doc.add_heading("11. Defects and external links", level=1)
doc.add_paragraph(
    "Replaces Osprey's Jira/QMetry TAM API. Defects are first-class "
    "ReportStack resources with summary, status, defect_type, and "
    "external_link. Use the external_link to point at a Jira ticket, GitHub "
    "issue, or any URL — ReportStack doesn't reach into your tracker, it just "
    "links out."
)
add_table(
    ["Status", "Meaning"],
    [
        ["OPEN", "Newly filed, not assigned"],
        ["IN_PROGRESS", "Owner picked it up"],
        ["FIXED", "Code merged; verify on next run"],
        ["WONT_FIX", "Acknowledged, not actioned"],
        ["DUPLICATE", "Reference an existing defect via external_link"],
    ],
    widths_in=[1.5, 5.0],
)

section_break()

# ============================================================
# 12. KPIs
# ============================================================
doc.add_heading("12. Performance KPIs", level=1)
doc.add_paragraph(
    "Radio-relevant KPIs supersede Osprey's TV-specific ones. Each KPI test "
    "measures a value, attaches it to the test item via user_properties, and "
    "asserts a threshold. The Trends widget plots the value across launches."
)
add_table(
    ["KPI", "Measured value", "Suggested threshold"],
    [
        ["Lock time at 433 MHz", "Seconds from tune() to is_locked()=true", "< 2.0 s"],
        ["Channel switch time", "Seconds between two tune() calls reaching lock", "< 1.0 s"],
        ["Power-on to first response", "Seconds from power_on to first responsive cmd", "< 5.0 s"],
        ["RSSI stability (median over 60 s)", "Standard deviation of dBm samples", "< 3 dB"],
        ["Bit error rate (when supported)", "Fraction of erroneous bits", "< 1e-5"],
        ["Reboot recovery time", "Seconds from reboot to lock at known frequency", "< 30 s"],
    ],
    widths_in=[2.3, 2.7, 1.5],
)
doc.add_heading("Reporting a KPI", level=2)
add_code(
    "@pytest.mark.kpi\n"
    "def test_lock_time_at_433mhz_under_2s(radio_up, request):\n"
    "    elapsed = tune_and_wait_for_lock(radio_up, 433.92, timeout_seconds=5.0)\n"
    "    request.node.user_properties.append(\n"
    "        (\"kpi:lock_time_433_seconds\", elapsed)\n"
    "    )\n"
    "    assert elapsed < 2.0"
)

section_break()

# ============================================================
# 13. ARTIFACTS
# ============================================================
doc.add_heading("13. Artifacts via ReportStack attachments", level=1)
doc.add_paragraph(
    "Replaces Nexus binary repos. The ReportStack backend stores attachments "
    "in MinIO (S3-compatible) and serves them by stable URL. From a test, "
    "register a path you want uploaded:"
)
add_code(
    "import pytest\n"
    "from pathlib import Path\n"
    "\n"
    "def test_with_recording(radio_up, request, tmp_path):\n"
    "    log_file = tmp_path / \"capture.log\"\n"
    "    log_file.write_text(\"...\")\n"
    "    request.node.user_properties.append((\"attachment\", str(log_file)))"
)
doc.add_paragraph(
    "Maximum 20 MB per attachment. For larger files (long captures, raw IQ "
    "data), upload directly to your own object store and set the test item's "
    "comment with the URL."
)

section_break()

# ============================================================
# 14. NOTIFICATIONS
# ============================================================
doc.add_heading("14. Notifications", level=1)
doc.add_paragraph(
    "ReportStack handles notification rules in the dashboard (Settings → "
    "Notifications). Configure them once per project — no test-side changes "
    "needed. Triggers include launch finished, launch failed, threshold "
    "exceeded; recipients are email, Slack, or Telegram."
)
add_callout(
    "Email Freemarker templates from Osprey don't translate. Use ReportStack's "
    "in-app notification settings instead.",
    "note",
)

section_break()

# ============================================================
# 15. JENKINS
# ============================================================
doc.add_heading("15. Jenkins CI/CD pipeline", level=1)
doc.add_paragraph(
    "Replace mvn integration-test with pytest. Everything else (cron triggers, "
    "agent labels, Nexus credentials) stays as it was."
)
add_code(
    "// Jenkinsfile (declarative)\n"
    "pipeline {\n"
    "  agent { label 'radio-lab' }\n"
    "  triggers { cron('H 1 * * *') }   // nightly 1 AM\n"
    "  environment {\n"
    "    AR_URL   = credentials('AR_URL')\n"
    "    AR_TOKEN = credentials('AR_TOKEN')\n"
    "  }\n"
    "  stages {\n"
    "    stage('Install') {\n"
    "      steps {\n"
    "        sh 'pip install -e test_framework_starter/'\n"
    "        sh 'pip install -e plugins/pytest-automation-reports/'\n"
    "      }\n"
    "    }\n"
    "    stage('Test') {\n"
    "      steps {\n"
    "        sh '''\n"
    "          pytest test_framework_starter/tests/ \\\n"
    "            --ar-url=$AR_URL \\\n"
    "            --ar-launch-name=\"$JOB_NAME #$BUILD_NUMBER\" \\\n"
    "            --ar-launch-description=\"build $GIT_COMMIT\" \\\n"
    "            --ar-tag=nightly --ar-auto-analyze \\\n"
    "            --device-host=10.0.0.42 \\\n"
    "            -m \"smoke or (regression and not slow)\"\n"
    "        '''\n"
    "      }\n"
    "    }\n"
    "  }\n"
    "}"
)

section_break()

# ============================================================
# 16. MIGRATION MAPPING
# ============================================================
doc.add_heading("16. Migration mapping: Osprey → this framework", level=1)
add_table(
    ["Osprey component", "ReportStack pytest equivalent", "Notes"],
    [
        ["Java 11", "Python 3.12", ""],
        ["Maven (pom.xml)", "pyproject.toml", "Optional groups replace profiles"],
        ["TestNG", "pytest", "Markers replace @Test groups; fixtures replace @Before/@After"],
        ["TestNG XML suites", "pytest -m expressions", "No more file-per-suite"],
        ["AspectJ AOP", "Pytest hooks + plugins", "First-class — no weaving needed"],
        ["OAT Core BaseTest", "Composable fixtures in conftest", "Avoid base-class inheritance"],
        ["Page Object Model", "RadioDevice + impls", "Same pattern, hardware-flavoured"],
        ["Locator JSON files", "Command catalog JSONs", "Same idea, different vocabulary"],
        ["DataProvider", "@pytest.mark.parametrize", ""],
        ["RetryListener", "pytest-rerunfailures", ""],
        ["FailListener", "pytest_runtest_makereport hook", ""],
        ["OspreyReportPortalListener", "pytest-automation-reports plugin", ""],
        ["AWS Lambda triage", "Ollama via ReportStack analyzer", "Local, no API key"],
        ["Jira/QMetry TAM API", "ReportStack defects + external_link", "Manual link out"],
        ["Confluence auto-update", "ReportStack Dashboards", "Configurable widgets"],
        ["Nexus artifact storage", "ReportStack attachments → MinIO", "20 MB cap"],
        ["Email Freemarker", "ReportStack notification rules", "In-app config"],
        ["Appium / UIAutomator2 / ADB", "paramiko (SSH) / pyserial", "Linux radios, not Android"],
        ["Charles Proxy HAR parsing", "(out of scope)", "Add as a later module if needed"],
    ],
    widths_in=[2.0, 2.5, 2.0],
)

section_break()

# ============================================================
# 17. WORKED EXAMPLE
# ============================================================
doc.add_heading("17. Worked example — one full test, end to end", level=1)
doc.add_paragraph(
    "A complete trace from CLI invocation to ReportStack dashboard. Assumes a "
    "radio at 10.0.0.42 with the wnc_radio.json catalog and the AR backend "
    "at http://reports.local:8000."
)
doc.add_heading("The invocation", level=2)
add_code(
    "pytest test_framework_starter/tests/test_kpi.py \\\n"
    "  --ar-url=http://reports.local:8000/api/v1 \\\n"
    "  --ar-launch-name=\"lock-time spot check\" \\\n"
    "  --ar-tag=manual --ar-auto-analyze \\\n"
    "  --device-host=10.0.0.42 --device-user=root \\\n"
    "  --device-cmds=test_framework_starter/commands/examples/wnc_radio.json \\\n"
    "  -m kpi"
)
doc.add_heading("What happens", level=2)
add_bullets([
    "1. The AR plugin POSTs to /launches/ and remembers the launch ID.",
    "2. Pytest collects the kpi-marked test (test_lock_time_at_433mhz_under_2s).",
    "3. The session-scoped radio fixture opens an SSH connection to 10.0.0.42 and validates the command catalog.",
    "4. The test-scoped radio_up fixture calls bring_up — power_on then a 1.5 s settle.",
    "5. The test calls tune_and_wait_for_lock(433.92, timeout_seconds=5.0); records 1.34 s.",
    "6. The test pushes (\"kpi:lock_time_433_seconds\", 1.34) onto user_properties; asserts < 2.0.",
    "7. The AR plugin reports the test result + the KPI metric.",
    "8. Teardown — radio_up calls shut_down, radio fixture closes the SSH session.",
    "9. Session end — the AR plugin PUTs /launches/{id}/finish with status PASSED, then POSTs /launches/{id}/analyze (because --ar-auto-analyze).",
    "10. In the dashboard, the launch shows PASSED; the Trends page plots 1.34 s for the kpi:lock_time_433_seconds metric.",
])

doc.add_heading("If it had failed", level=2)
add_bullets([
    "The test would raise AssertionError with the elapsed time in its message.",
    "pytest_runtest_makereport stamps rep_call=failed on the item.",
    "attach_diagnostics_on_fail (autouse fixture) snapshots lock state, RSSI, and temperature into a JSON file.",
    "The AR plugin uploads the JSON as an attachment.",
    "After the launch finishes, the analyzer triggers — Ollama classifies the failure (likely PRODUCT_BUG or AUTOMATION_BUG given the message).",
    "On the Test Detail page, the triager sees the assertion, the diagnostic snapshot, and the AI's call. They can override via the Defect Selector if they disagree.",
])

doc.add_paragraph()
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = final.add_run(
    f"Generated {datetime.date.today().strftime('%B %d, %Y')}  ·  "
    "Companion to test_framework_starter/"
)
r.font.size = Pt(9); r.font.color.rgb = BRAND_LIGHT


output_path = "/Users/rashadjafarzadeh/Projects/automation-reports/ReportStack_pytest_Framework_Guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
