"""Generate ReportStack project documentation as Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for row in rows:
        r = t.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = str(val)
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

# ========================
# TITLE PAGE
# ========================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ReportStack")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Self-Hosted Test Automation Reporting Platform")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f"Project Documentation\nVersion 1.0\n{datetime.date.today().strftime('%B %d, %Y')}")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ========================
# TABLE OF CONTENTS
# ========================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Project Overview",
    "3. Architecture & Tech Stack",
    "4. System Components",
    "5. Data Model",
    "6. API Reference",
    "7. Frontend Pages & Features",
    "8. AI Failure Analysis",
    "9. Telegram CEO Bot",
    "10. pytest Plugin",
    "11. Storage & Infrastructure",
    "12. Completed Work",
    "13. Planned Work & Roadmap",
    "14. Deployment Guide",
    "15. Configuration Reference",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ========================
# 1. EXECUTIVE SUMMARY
# ========================
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "ReportStack is a self-hosted test automation reporting platform designed for QA/SDET teams "
    "running Selenium, Playwright, and other test frameworks. It provides centralized test result "
    "collection, visualization, and AI-powered failure classification — all running on-premise "
    "with zero external API dependencies."
)
doc.add_paragraph(
    "The platform is comparable to ReportPortal.io but with a modern UI, local AI analysis via "
    "Ollama, and a Telegram bot for executive reporting. It ingests test results from CI/CD pipelines "
    "through a pytest plugin and presents them through an interactive React dashboard."
)

doc.add_heading("Key Capabilities", level=2)
bullets = [
    "Real-time test result ingestion via pytest plugin with automatic screenshot capture",
    "Interactive dashboards with customizable widget grids (charts, tables, metrics)",
    "AI-powered failure classification using local Ollama LLM (no cloud dependency)",
    "Cross-launch test history tracking and trend analysis",
    "Defect management with ReportPortal-style defect type selector",
    "Team member management with role-based access (Admin/Manager/Member/Viewer)",
    "Telegram bot for CEO/stakeholder reporting with natural language queries",
    "S3-compatible attachment storage via MinIO",
    "Full REST API for integration with any test framework or CI/CD system",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

# ========================
# 2. PROJECT OVERVIEW
# ========================
doc.add_page_break()
doc.add_heading("2. Project Overview", level=1)

doc.add_heading("Problem Statement", level=2)
doc.add_paragraph(
    "QA teams running large test suites lack a centralized, self-hosted platform to collect results, "
    "analyze failures, and report to stakeholders. Existing solutions like ReportPortal are complex "
    "to deploy and maintain. Cloud-based solutions raise data privacy concerns for enterprise teams."
)

doc.add_heading("Solution", level=2)
doc.add_paragraph(
    "ReportStack provides a lightweight, Docker-based reporting platform that any team can deploy "
    "in minutes. It uses local AI for failure analysis, stores all data on-premise, and offers "
    "a modern React UI with real-time updates."
)

doc.add_heading("Target Users", level=2)
bullets = [
    "QA Engineers & SDETs — View test results, analyze failures, manage defects",
    "QA Leads & Managers — Track trends, monitor pass rates, manage team members",
    "Engineering Directors & CEOs — Get high-level reports via Telegram bot",
    "DevOps Engineers — Deploy and configure the platform via Docker Compose",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

# ========================
# 3. ARCHITECTURE & TECH STACK
# ========================
doc.add_page_break()
doc.add_heading("3. Architecture & Tech Stack", level=1)

doc.add_heading("Technology Stack", level=2)
add_table(
    ["Layer", "Technology", "Version"],
    [
        ["Backend API", "FastAPI + SQLAlchemy + Alembic", "Python 3.12, FastAPI 0.111"],
        ["Database", "PostgreSQL", "16 (Alpine)"],
        ["Frontend", "React + TypeScript", "React 19, TypeScript 4.9"],
        ["Charts", "Recharts", "3.8"],
        ["AI Engine", "Ollama (local LLM)", "mistral:7b default"],
        ["Object Storage", "MinIO (S3-compatible)", "Latest"],
        ["Test Plugin", "pytest plugin", "pytest >= 7.0"],
        ["Bot", "python-telegram-bot + Claude API", "Latest"],
        ["Infrastructure", "Docker Compose", "v3.8"],
    ],
)

doc.add_heading("System Architecture", level=2)
doc.add_paragraph(
    "The system follows a microservices architecture deployed via Docker Compose. "
    "All services communicate over a Docker bridge network:"
)
components = [
    "pytest Plugin → sends test results to Backend API via HTTP",
    "FastAPI Backend (:8000) → REST API, connects to PostgreSQL, Ollama, and MinIO",
    "PostgreSQL (:5432) → persistent data storage for all entities",
    "MinIO (:9000/:9001) → S3-compatible object storage for attachments (screenshots, logs, videos)",
    "Ollama (:11434) → local LLM for AI failure classification",
    "React Frontend (:3000) → SPA served via nginx, proxies API calls to backend",
    "Telegram Bot → polls Telegram API, uses Claude for reasoning, queries backend for live data",
]
for c in components:
    doc.add_paragraph(c, style="List Bullet")

# ========================
# 4. SYSTEM COMPONENTS
# ========================
doc.add_page_break()
doc.add_heading("4. System Components", level=1)

doc.add_heading("4.1 Backend (FastAPI)", level=2)
doc.add_paragraph(
    "The backend is a Python FastAPI application providing a RESTful API. It handles all business "
    "logic including test result ingestion, attachment management, AI analysis orchestration, "
    "and user/project management."
)
doc.add_heading("Project Structure", level=3)
structure = [
    "app/api/ — FastAPI routers (launches, test_items, logs, attachments, analyses, comments, defects, members, project_settings, dashboards, test_history)",
    "app/models/ — SQLAlchemy ORM models",
    "app/schemas/ — Pydantic request/response schemas",
    "app/services/ — Business logic (ai_analyzer.py, storage.py)",
    "app/database.py — Engine, session factory, Base class",
    "app/main.py — Application setup, CORS, router registration, startup events",
    "migrations/ — Alembic database migrations",
]
for s in structure:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("4.2 Frontend (React)", level=2)
doc.add_paragraph(
    "The frontend is a React 19 SPA with TypeScript, using Recharts for data visualization. "
    "It follows a design system with CSS custom properties and reusable component classes."
)

doc.add_heading("4.3 MinIO Object Storage", level=2)
doc.add_paragraph(
    "Attachments (screenshots, log files, videos) are stored in MinIO, an S3-compatible object "
    "storage system. The backend uses the minio Python SDK to upload, download, and delete files. "
    "Files are organized as {launch_id}/{item_id}/{uuid}_{filename} in the 'attachments' bucket. "
    "The bucket is auto-created on backend startup."
)

doc.add_heading("4.4 PostgreSQL Database", level=2)
doc.add_paragraph(
    "All structured data is stored in PostgreSQL 16. The schema is managed via Alembic migrations. "
    "The database stores launches, test items, logs, attachments metadata, analyses, comments, "
    "defects, members, project settings, dashboards, and widgets."
)

# ========================
# 5. DATA MODEL
# ========================
doc.add_page_break()
doc.add_heading("5. Data Model", level=1)

doc.add_heading("Entity Relationships", level=2)
doc.add_paragraph(
    "Launch (1) → (N) TestItem → (N) TestLog, Attachment, FailureAnalysis, Comment, Defect. "
    "Launch also has direct relationships to Attachment (launch-level) and Comment (launch-level). "
    "Member and ProjectSettings are standalone entities. Dashboard (1) → (N) Widget."
)

doc.add_heading("Core Entities", level=2)
add_table(
    ["Entity", "Key Fields", "Description"],
    [
        ["Launch", "name, status, start_time, end_time, total, passed, failed, skipped", "A test execution run (e.g., CI build)"],
        ["TestItem", "name, suite, status, duration_ms, error_message, stack_trace", "Individual test case result"],
        ["TestLog", "timestamp, level, message, step_name, order_index", "Log entry within a test"],
        ["Attachment", "file_name, file_path, file_size, content_type, attachment_type", "File stored in MinIO"],
        ["FailureAnalysis", "defect_type, confidence, reasoning, source, model_name", "AI or manual failure classification"],
        ["Comment", "author, text, test_item_id/launch_id", "User comment on test or launch"],
        ["Defect", "summary, description, status, external_id, external_url", "Tracked defect linked to test"],
        ["Member", "name, email, role", "Team member with role-based access"],
        ["ProjectSettings", "project_name, auto_analysis_enabled, retention_days, ...", "Singleton project configuration"],
        ["Dashboard", "name, description, owner", "Custom dashboard container"],
        ["Widget", "widget_type, title, config, width, order_index", "Dashboard widget"],
    ],
)

doc.add_heading("Enumerations", level=2)
add_table(
    ["Enum", "Values"],
    [
        ["LaunchStatus", "IN_PROGRESS, PASSED, FAILED, STOPPED"],
        ["TestStatus", "PASSED, FAILED, SKIPPED, ERROR"],
        ["LogLevel", "TRACE, DEBUG, INFO, WARN, ERROR"],
        ["AttachmentType", "SCREENSHOT, LOG_FILE, VIDEO, OTHER"],
        ["DefectType", "PRODUCT_BUG, AUTOMATION_BUG, SYSTEM_ISSUE, NO_DEFECT, TO_INVESTIGATE"],
        ["AnalysisSource", "AI_AUTO, MANUAL"],
        ["DefectStatus", "OPEN, IN_PROGRESS, FIXED, WONT_FIX, DUPLICATE"],
        ["MemberRole", "ADMIN, MANAGER, MEMBER, VIEWER"],
    ],
)

# ========================
# 6. API REFERENCE
# ========================
doc.add_page_break()
doc.add_heading("6. API Reference", level=1)
doc.add_paragraph("All endpoints are under /api/v1. The API follows RESTful conventions.")

sections = [
    ("Launches", [
        ["POST", "/launches/", "Create a new launch"],
        ["GET", "/launches/", "List launches (paginated)"],
        ["GET", "/launches/{id}", "Get single launch"],
        ["PUT", "/launches/{id}/finish", "Finish launch with status"],
        ["DELETE", "/launches/{id}", "Delete launch and all children"],
    ]),
    ("Test Items", [
        ["POST", "/launches/{id}/items/", "Create single test item"],
        ["POST", "/launches/{id}/items/batch", "Create batch of test items"],
        ["GET", "/launches/{id}/items/", "List items (filter by status, suite)"],
        ["GET", "/launches/{id}/items/{item_id}", "Get single test item"],
    ]),
    ("Logs", [
        ["POST", "/launches/{id}/items/{item_id}/logs/", "Create single log"],
        ["POST", "/launches/{id}/items/{item_id}/logs/batch", "Create batch of logs"],
        ["GET", "/launches/{id}/items/{item_id}/logs/", "List logs (filter by level)"],
    ]),
    ("Attachments", [
        ["POST", "/launches/{id}/items/{item_id}/attachments", "Upload file (multipart, 20MB max)"],
        ["POST", "/launches/{id}/attachments", "Upload launch-level attachment"],
        ["GET", "/launches/{id}/items/{item_id}/attachments", "List item attachments"],
        ["GET", "/attachments/{id}/file", "Serve/download file from MinIO"],
        ["DELETE", "/attachments/{id}", "Delete file from MinIO + DB"],
    ]),
    ("AI Analysis", [
        ["POST", "/launches/{id}/analyze", "Trigger analysis for all failures (async)"],
        ["POST", "/launches/{id}/items/{item_id}/analyze", "Trigger single item analysis"],
        ["GET", "/launches/{id}/items/{item_id}/analyses", "Get analyses for item"],
        ["GET", "/launches/{id}/analysis-summary", "Aggregate defect type counts"],
        ["PUT", "/launches/{id}/items/{item_id}/analyses/{a_id}", "Manual override"],
    ]),
    ("Comments", [
        ["POST", "/launches/{id}/items/{item_id}/comments/", "Create item comment"],
        ["GET", "/launches/{id}/items/{item_id}/comments/", "List item comments"],
        ["POST", "/launches/{id}/comments/", "Create launch comment"],
        ["GET", "/launches/{id}/comments/", "List launch comments"],
        ["PUT", "/comments/{id}", "Update comment"],
        ["DELETE", "/comments/{id}", "Delete comment"],
    ]),
    ("Defects", [
        ["POST", "/launches/{id}/items/{item_id}/defects/", "Create defect"],
        ["GET", "/launches/{id}/items/{item_id}/defects/", "List defects"],
        ["PUT", "/defects/{id}", "Update defect"],
        ["DELETE", "/defects/{id}", "Delete defect"],
    ]),
    ("Members", [
        ["GET", "/members/", "List all members"],
        ["POST", "/members/", "Add member (unique email)"],
        ["PUT", "/members/{id}", "Update member role"],
        ["DELETE", "/members/{id}", "Remove member"],
    ]),
    ("Project Settings", [
        ["GET", "/settings/", "Get settings (auto-creates defaults)"],
        ["PUT", "/settings/", "Update settings"],
    ]),
    ("Dashboards", [
        ["POST", "/dashboards/", "Create dashboard"],
        ["GET", "/dashboards/", "List dashboards"],
        ["GET", "/dashboards/{id}", "Get single dashboard"],
        ["PUT", "/dashboards/{id}", "Update dashboard"],
        ["DELETE", "/dashboards/{id}", "Delete dashboard"],
        ["POST", "/dashboards/{id}/widgets/", "Add widget"],
        ["DELETE", "/dashboards/{id}/widgets/{w_id}", "Remove widget"],
    ]),
    ("Test History", [
        ["GET", "/items/history?name=...&limit=20", "Cross-launch test history"],
        ["GET", "/items/most-failed?limit=50", "Top N most-failed test cases"],
    ]),
    ("Health", [
        ["GET", "/health", "Health check"],
    ]),
]

for section_name, endpoints in sections:
    doc.add_heading(section_name, level=2)
    add_table(["Method", "Endpoint", "Description"], endpoints)

# ========================
# 7. FRONTEND PAGES
# ========================
doc.add_page_break()
doc.add_heading("7. Frontend Pages & Features", level=1)

pages = [
    ("Dashboard (/)", "Overview metrics dashboard with key statistics across all launches."),
    ("Launches (/launches)", "Paginated launch list with status badges, pass rate bars, and real-time polling for IN_PROGRESS launches. Shows LIVE badge with pulse animation for active launches."),
    ("Launch Detail (/launches/:id)", "Two-tab view (Tests + Comments). Tests tab shows test items in a table with click-to-navigate to TestDetail. Failed tests show inline DefectSelector right rail. Real-time polling every 5 seconds for active launches."),
    ("Test Detail (/launches/:id/items/:itemId)", "Full test detail page with error message, stack trace (white background), log viewer with inline screenshot thumbnails, attachment gallery, AI analysis panel, comments section, and cross-launch HistoryStrip component."),
    ("Dashboards (/dashboards)", "Customizable dashboard page with dropdown selector to switch between dashboards. 9 widget types: Launch Statistics Area/Bar charts, Overall Statistics Panel/Donut, Launch Duration chart, Failed Cases Trend, Passing Rate Pie, Launch Table, Most Failed Test Cases table. Add/remove widgets dynamically."),
    ("Trends (/trends)", "Launch trend analysis with pass rate area chart, test results stacked bar chart, duration line chart, and launch history table. Filters by name and limit (10/20/50/all)."),
    ("Members (/members)", "Team member management table with role assignment (Admin/Manager/Member/Viewer)."),
    ("Settings (/settings)", "9-tab settings page: General, Integrations, Notifications, Defect Types, Log Types, Analyzer, Pattern-analysis, Demo Data, Quality Gates."),
    ("Profile (/profile)", "User profile page."),
]

for name, desc in pages:
    doc.add_heading(name, level=2)
    doc.add_paragraph(desc)

doc.add_heading("Dashboard Widget Types", level=2)
add_table(
    ["Widget Type", "Visualization", "Data Source"],
    [
        ["LAUNCH_STATS_AREA", "Stacked area chart (passed/failed/skipped)", "Launch statistics"],
        ["LAUNCH_STATS_BAR", "Stacked bar chart", "Launch statistics"],
        ["OVERALL_STATS_PANEL", "Big number metrics panel", "Aggregated totals"],
        ["OVERALL_STATS_DONUT", "Pie/donut chart by status", "Aggregated totals"],
        ["LAUNCH_DURATION", "Line chart of durations", "Launch start/end times"],
        ["FAILED_TREND", "Area chart of failure counts", "Launch failure counts"],
        ["PASS_RATE_PIE", "Pie chart (passed vs not passed)", "Aggregated totals"],
        ["LAUNCH_TABLE", "Clickable table with status badges", "Recent launches"],
        ["MOST_FAILED", "Table with failure rate bars", "Cross-launch aggregation"],
    ],
)

doc.add_heading("Design System", level=2)
doc.add_paragraph(
    "The frontend uses a comprehensive CSS design system with custom properties (CSS variables) "
    "for colors, spacing, typography, shadows, and border radius. All components use semantic "
    "class names defined in components.css and extras.css. Key design tokens include semantic "
    "colors (--color-passed, --color-failed, etc.), a spacing scale, Inter/JetBrains Mono fonts, "
    "and consistent shadow/radius values."
)

doc.add_heading("Branding & Visual Identity", level=2)
doc.add_paragraph(
    "ReportStack's visual identity centers on a stacked-bars mark — three rounded horizontal "
    "bars in three blue tones that read as both a bar chart and a stack of reports. The wordmark "
    "renders \"Report\" in slate (#0F172A) and \"Stack\" in primary blue (#1E40AF), reinforcing "
    "the brand with the same color split that runs through the icon."
)

doc.add_heading("Brand color palette", level=3)
add_table(
    ["Token", "Hex", "Use"],
    [
        ["--color-primary", "#1E40AF", "Default CTA, links, breadcrumb, active nav"],
        ["--color-primary-hover", "#1E3A8A", "Hover state for primary"],
        ["--color-primary-active", "#3B82F6", "Bright accent — sidebar stripe, focus ring"],
        ["--color-primary-soft", "#EFF6FF", "Hover row tints, badge backgrounds"],
        ["--color-primary-light", "#DBEAFE", "Soft surfaces, subtle borders"],
        ["--gradient-brand", "#1E40AF → #0B2E7C", "Logo tile, login top stripe"],
        ["--gradient-brand-soft", "#EFF6FF → #DBEAFE", "Reserved for hero / empty states"],
    ],
)

doc.add_heading("Logo assets", level=3)
add_table(
    ["Asset", "Path", "Purpose"],
    [
        ["Horizontal lockup", "frontend/src/logo.svg", "Imported by Sidebar/Login. Text vectorized to paths so it renders identically without a font dependency."],
        ["Centered icon mark", "Inline in App.tsx", "3-bar SVG inside the 36×36 brand-gradient tile in the sidebar."],
        ["Favicon", "frontend/public/favicon.ico", "Multi-size (16/32/48/64/128/256) bundled from the rounded blue tile."],
        ["PWA icons", "frontend/public/logo192.png, logo512.png", "Apple touch icon and PWA install icon."],
    ],
)
doc.add_paragraph(
    "PWA metadata in manifest.json (name, short_name, theme_color #1E40AF) and the title and "
    "theme-color meta in index.html are aligned with the brand. The pre-existing Create React "
    "App placeholders have been replaced."
)

doc.add_heading("Surface treatments", level=3)
surface_bullets = [
    "Sidebar: 36×36 logo tile uses --gradient-brand with an inset highlight; active nav link shows a 3px --color-primary-active left stripe (via ::before) over a soft blue tint.",
    "Login: layered radial brand-tinted gradient backdrop, 16-radius card with --shadow-xl, and a 3px --gradient-brand stripe across the top edge of the card.",
    "Launches list: page-title uses -0.015em letter-spacing; metric cards lift on hover with a brand-tinted shadow; in-progress badges adopt brand blue.",
    "Launch detail: back button hover uses --color-primary-soft + --color-primary-light border for a subtle brand accent.",
    "Buttons: btn-primary hover gains a brand-tinted shadow (rgba(30,64,175,0.32)) instead of the previous neutral one.",
]
for b in surface_bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Branding rollout status", level=3)
doc.add_paragraph(
    "Wave 1 (May 2026) is live on main as two commits: \"Add ReportStack logo: replace CRA "
    "placeholders with new mark\" and \"Apply ReportStack brand to UI: tokens + key surfaces\". "
    "Wave 2 — Settings ps-* token sweep, Members, Dashboards, LaunchDetail polish, empty/loading "
    "states, and a future dark-mode pass — is queued and documented in HANDOFF_BRANDING.md at the "
    "repo root, with a suggested per-page PR order so each diff stays reviewable."
)

# ========================
# 8. AI FAILURE ANALYSIS
# ========================
doc.add_page_break()
doc.add_heading("8. AI Failure Analysis", level=1)

doc.add_heading("How It Works", level=2)
steps = [
    "POST /launches/{id}/analyze triggers BackgroundTasks for all FAILED/ERROR test items",
    "For each item, the system gathers: test name, suite, error_message, stack_trace (last 50 lines), recent logs (last 20)",
    "Sends structured prompt to Ollama with classification instructions",
    "Expects JSON response: {defect_type, confidence (0-1), reasoning}",
    "Confidence < 0.4 automatically flags as TO_INVESTIGATE",
    "Retries once on JSON parse failure; falls back to TO_INVESTIGATE if Ollama is unreachable",
    "Results stored in FailureAnalysis table with source=AI_AUTO",
    "Users can manually override via PUT endpoint (source becomes MANUAL)",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("Defect Type Classification", level=2)
add_table(
    ["Defect Type", "Description", "Examples"],
    [
        ["PRODUCT_BUG", "Actual bug in the application under test", "Unexpected error, wrong behavior"],
        ["AUTOMATION_BUG", "Bug in the test code itself", "Stale selector, wrong assertion"],
        ["SYSTEM_ISSUE", "Infrastructure/environment problem", "Timeout, connection refused, OOM"],
        ["NO_DEFECT", "Test is correctly failing (expected)", "Feature not yet implemented"],
        ["TO_INVESTIGATE", "Requires human review", "Low confidence, ambiguous failure"],
    ],
)

doc.add_heading("Configuration", level=2)
add_table(
    ["Variable", "Default", "Description"],
    [
        ["OLLAMA_BASE_URL", "http://localhost:11434", "Ollama API endpoint"],
        ["OLLAMA_MODEL", "mistral:7b", "LLM model for analysis"],
    ],
)

# ========================
# 9. TELEGRAM BOT
# ========================
doc.add_page_break()
doc.add_heading("9. Telegram CEO Bot", level=1)

doc.add_paragraph(
    "The Telegram bot serves as an intelligent assistant powered by Claude AI. It carries full "
    "project knowledge and can query the live system for real-time data. Designed for CEO/stakeholder "
    "reporting without requiring access to the web dashboard."
)

doc.add_heading("Commands", level=2)
add_table(
    ["Command", "Description"],
    [
        ["/start, /help", "Welcome message and command list"],
        ["/status", "Backend health check + launch count"],
        ["/launches", "Recent 8 launches with pass rates"],
        ["/launch <id>", "Detailed launch info + analysis summary"],
        ["/failures <id>", "List failed tests with error previews"],
        ["/analyze <id>", "Trigger AI failure analysis"],
        ["/team", "List team members"],
        ["/clear", "Reset conversation history"],
        ["Free text", "Routed to Claude agent with project context"],
    ],
)

doc.add_heading("Claude Agent Tools", level=2)
tools = [
    "check_system_health — GET /health",
    "get_launches — GET /launches/ (paginated)",
    "get_launch_detail — GET /launches/{id}",
    "get_failed_tests — GET /items/ filtered by FAILED + ERROR",
    "get_analysis_summary — GET /analysis-summary",
    "trigger_analysis — POST /launches/{id}/analyze",
    "get_test_history — GET /items/history",
    "get_test_logs — GET /items/{id}/logs",
    "get_item_defects — GET /items/{id}/defects",
    "get_dashboards — GET /dashboards/",
    "get_members — GET /members/",
    "get_settings — GET /settings/",
]
for t in tools:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("Configuration", level=2)
add_table(
    ["Variable", "Description"],
    [
        ["TELEGRAM_BOT_TOKEN", "Token from @BotFather"],
        ["ANTHROPIC_API_KEY", "Claude API key"],
        ["CLAUDE_MODEL", "Default: claude-sonnet-4-6"],
        ["AR_BACKEND_URL", "Backend API URL (default: http://backend:8000/api/v1)"],
        ["AR_FRONTEND_URL", "Frontend URL for generating dashboard links"],
        ["ALLOWED_USER_IDS", "Comma-separated Telegram user IDs (empty = allow all)"],
        ["MAX_HISTORY", "Conversation turns to retain per user (default 20)"],
    ],
)

# ========================
# 10. PYTEST PLUGIN
# ========================
doc.add_page_break()
doc.add_heading("10. pytest Plugin", level=1)

doc.add_paragraph(
    "The pytest plugin automatically collects test results and sends them to the ReportStack backend. "
    "It captures test names, suites, statuses, durations, errors, stack traces, logs, and screenshots."
)

doc.add_heading("Installation & Usage", level=2)
doc.add_paragraph("pip install -e plugins/pytest-automation-reports/", style="No Spacing")
doc.add_paragraph()
doc.add_paragraph(
    'pytest --ar-url=http://server:8000/api/v1 --ar-launch-name="Regression Run" '
    '--ar-launch-description="Build #123" --ar-auto-analyze',
    style="No Spacing",
)

doc.add_heading("Screenshot Fixture", level=2)
doc.add_paragraph(
    "The plugin provides a report_screenshot fixture for capturing Selenium/Playwright screenshots "
    "during tests. Screenshots are queued and uploaded as attachments after test completion."
)

doc.add_heading("Plugin Components", level=2)
add_table(
    ["File", "Purpose"],
    [
        ["plugin.py", "pytest hooks (session start/finish, test report) + fixture registration"],
        ["client.py", "HTTP client for backend API communication"],
        ["config.py", "CLI options (--ar-url, --ar-launch-name, etc.)"],
        ["log_capture.py", "logging.Handler for capturing test logs"],
        ["screenshot.py", "Selenium/Playwright screenshot capture helper"],
    ],
)

# ========================
# 11. STORAGE & INFRASTRUCTURE
# ========================
doc.add_page_break()
doc.add_heading("11. Storage & Infrastructure", level=1)

doc.add_heading("MinIO Object Storage", level=2)
doc.add_paragraph(
    "Attachments are stored in MinIO, an S3-compatible object storage system. The backend uses "
    "the minio Python SDK for all file operations. Files are organized by launch and test item ID."
)
add_table(
    ["Operation", "Method", "Description"],
    [
        ["Upload", "storage.upload_file(key, data, content_type)", "Stores file in MinIO bucket"],
        ["Download", "storage.download_file(key)", "Returns (bytes, content_type)"],
        ["Delete", "storage.delete_file(key)", "Removes file from bucket"],
        ["Presigned URL", "storage.get_presigned_url(key, hours)", "Generates temporary direct download URL"],
        ["Ensure Bucket", "storage.ensure_bucket()", "Creates bucket if not exists (called on startup)"],
    ],
)

doc.add_heading("Docker Compose Services", level=2)
add_table(
    ["Service", "Image", "Ports", "Volumes"],
    [
        ["db", "postgres:16-alpine", "5432", "pgdata:/var/lib/postgresql/data"],
        ["minio", "minio/minio:latest", "9000 (API), 9001 (Console)", "minio_data:/data"],
        ["backend", "Custom (./backend)", "8000", "None"],
        ["frontend", "Custom (./frontend)", "3000 → 80 (nginx)", "None"],
        ["bot", "Custom (./bot)", "None (polling)", "None"],
        ["ollama", "ollama/ollama:latest", "11434", "ollama_data:/root/.ollama"],
    ],
)

# ========================
# 12. COMPLETED WORK
# ========================
doc.add_page_break()
doc.add_heading("12. Completed Work", level=1)

completed = [
    ("PR #1 — UI/UX Redesign", [
        "Complete design system with CSS custom properties",
        "Dark sidebar navigation with project sections",
        "Dashboard overview page with metric cards",
        "Launch list with status badges and pass rate bars",
        "Launch detail with two-tab layout (Tests + Comments)",
        "Test detail page with log viewer, screenshot gallery, analysis panel",
        "Members management page with role assignment",
        "9-tab Settings page (ReportPortal-style)",
        "Profile page",
    ]),
    ("PR #2 — Test Navigation & UI Fixes", [
        "Click-to-navigate from test rows to TestDetail page",
        "Stack trace display changed to white background with black text",
        "Inline screenshot thumbnails next to log entries",
    ]),
    ("PR #3 — Dashboards, HistoryStrip, DefectSelector", [
        "Dashboard CRUD page with card grid",
        "HistoryStrip component showing cross-launch test history with colored status bars",
        "Inline DefectSelector right rail in LaunchDetail for failed tests",
        "Backend endpoints for dashboards (CRUD + widgets) and test history",
    ]),
    ("PR #4 — Telegram Bot Update", [
        "Updated system prompt with all current pages and APIs",
        "Added 7 new tools (test history, logs, defects, dashboards, members, settings)",
        "Added /team command",
    ]),
    ("PR #5 — Trends Page & Real-time Polling", [
        "Launch Trends page with pass rate area chart, test results bar chart, duration line chart",
        "Real-time polling for IN_PROGRESS launches (5-second interval)",
        "LIVE badge with pulse animation for active launches",
    ]),
    ("PR #6 — MinIO Storage & Dashboard Widgets", [
        "Switched attachment storage from local disk to MinIO (S3-compatible)",
        "Added MinIO service to Docker Compose",
        "Created storage service with upload/download/delete/presigned URL helpers",
        "Built dashboard view with 9 widget types (charts, tables, metrics)",
        "Added backend endpoint for most-failed test cases aggregation",
    ]),
    ("PR #7 — Merge Dashboards into Single Page", [
        "Combined dashboard list and detail view into single page",
        "Dropdown selector to switch between dashboards",
        "Widget grid renders inline — no navigation needed",
    ]),
]

for pr_name, items in completed:
    doc.add_heading(pr_name, level=2)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

# ========================
# 13. PLANNED WORK
# ========================
doc.add_page_break()
doc.add_heading("13. Planned Work & Roadmap", level=1)

doc.add_heading("Phase 1 — Core Enhancements", level=2)
phase1 = [
    "Authentication & Authorization — JWT-based auth with role enforcement across API endpoints",
    "Test Item Filtering — Advanced filtering by suite, status, duration, date range on launch detail",
    "Bulk Operations — Bulk defect assignment, bulk re-analysis, bulk status update",
    "Export — Export launch reports as PDF/HTML, export test data as CSV",
    "Retry Tracking — Link test retries together, show retry history on test detail",
]
for p in phase1:
    doc.add_paragraph(p, style="List Bullet")

doc.add_heading("Phase 2 — Advanced Features", level=2)
phase2 = [
    "Quality Gates — Define pass/fail criteria for launches (min pass rate, max failures, etc.)",
    "Pattern Analysis — Auto-group similar failures using AI embeddings",
    "Notifications — Email/Slack/Telegram alerts on launch completion, failure spikes",
    "Integration Cards — Jira, Azure DevOps ticket creation from defects",
    "Comparison View — Side-by-side comparison of two launches",
    "Flaky Test Detection — Auto-detect and flag flaky tests based on history patterns",
]
for p in phase2:
    doc.add_paragraph(p, style="List Bullet")

doc.add_heading("Phase 3 — Scale & Polish", level=2)
phase3 = [
    "Multi-project Support — Multiple projects with separate data isolation",
    "LDAP/SSO Integration — Enterprise authentication",
    "Dashboard Sharing — Public/private dashboards with share links",
    "Widget Drag & Drop — Rearrange dashboard widgets by dragging",
    "Custom Widget Sizing — Configurable widget width (1/2, 1/3, full)",
    "Dark Mode — Full dark theme support",
    "Mobile Responsive — Optimized layout for mobile devices",
    "Plugin SDK — Support for JUnit XML, Allure, and other test frameworks beyond pytest",
]
for p in phase3:
    doc.add_paragraph(p, style="List Bullet")

# ========================
# 14. DEPLOYMENT GUIDE
# ========================
doc.add_page_break()
doc.add_heading("14. Deployment Guide", level=1)

doc.add_heading("Quick Start", level=2)
steps = [
    "Clone the repository: git clone <repo-url>",
    "Copy environment file: cp .env.example .env",
    "Edit .env with your configuration (Telegram token, Anthropic key, etc.)",
    "Start all services: docker compose up -d --build",
    "Pull the AI model: docker compose exec ollama ollama pull mistral:7b",
    "Optional — seed demo data: python3 backend/seed_data.py",
    "Access the dashboard at http://localhost:3000",
    "Access MinIO console at http://localhost:9001 (minioadmin/minioadmin)",
    "Access API docs at http://localhost:8000/docs",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("Production Deployment (Linux Server)", level=2)
doc.add_paragraph(
    "For production, consider using bind mounts instead of named Docker volumes for easier "
    "backup and monitoring. Map data directories to known paths on the host:"
)
prod_mounts = [
    "/srv/reportstack/postgres → PostgreSQL data",
    "/srv/reportstack/minio → MinIO attachment storage",
    "/srv/reportstack/ollama → Ollama model files",
]
for m in prod_mounts:
    doc.add_paragraph(m, style="List Bullet")

doc.add_paragraph(
    "Set up regular backups (cron + rsync/tar) for /srv/reportstack/. "
    "Use a reverse proxy (nginx/Traefik) with SSL for production access."
)

# ========================
# 15. CONFIGURATION REFERENCE
# ========================
doc.add_page_break()
doc.add_heading("15. Configuration Reference", level=1)

doc.add_heading("Backend Environment Variables", level=2)
add_table(
    ["Variable", "Default", "Description"],
    [
        ["DATABASE_URL", "postgresql://postgres:postgres@db:5432/automation_reports", "PostgreSQL connection string"],
        ["OLLAMA_BASE_URL", "http://ollama:11434", "Ollama API endpoint"],
        ["OLLAMA_MODEL", "mistral:7b", "LLM model name"],
        ["MINIO_ENDPOINT", "minio:9000", "MinIO API endpoint"],
        ["MINIO_ACCESS_KEY", "minioadmin", "MinIO access key"],
        ["MINIO_SECRET_KEY", "minioadmin", "MinIO secret key"],
        ["MINIO_BUCKET", "attachments", "MinIO bucket name"],
        ["MINIO_SECURE", "false", "Use HTTPS for MinIO"],
    ],
)

doc.add_heading("Bot Environment Variables", level=2)
add_table(
    ["Variable", "Default", "Description"],
    [
        ["TELEGRAM_BOT_TOKEN", "(required)", "Telegram bot token from @BotFather"],
        ["ANTHROPIC_API_KEY", "(required)", "Claude API key"],
        ["CLAUDE_MODEL", "claude-sonnet-4-6", "Claude model to use"],
        ["AR_BACKEND_URL", "http://backend:8000/api/v1", "Backend API URL"],
        ["AR_FRONTEND_URL", "http://localhost:3000", "Frontend URL for links"],
        ["ALLOWED_USER_IDS", "(empty = all)", "Comma-separated allowed Telegram user IDs"],
        ["MAX_HISTORY", "20", "Conversation turns to retain"],
    ],
)

doc.add_heading("Frontend Environment Variables", level=2)
add_table(
    ["Variable", "Default", "Description"],
    [
        ["REACT_APP_API_URL", "/api/v1", "Backend API base URL (proxied via nginx)"],
    ],
)

# ========================
# FOOTER
# ========================
doc.add_page_break()
doc.add_heading("Document History", level=1)
add_table(
    ["Version", "Date", "Author", "Changes"],
    [
        ["1.0", datetime.date.today().strftime("%Y-%m-%d"), "ReportStack Team", "Initial comprehensive documentation"],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Document —")
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run.font.size = Pt(12)

# Save
output_path = "/Users/rashadjafarzadeh/Projects/automation-reports/ReportStack_Documentation.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
