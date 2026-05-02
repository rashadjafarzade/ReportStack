"""Generate ReportStack User Guide as a Word document.

Audience: QA engineers, QA leads, engineering stakeholders, DevOps engineers.
Companion to ReportStack_Documentation.docx (which is the system reference).
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

BRAND_DEEP = RGBColor(0x1E, 0x40, 0xAF)      # primary brand blue
BRAND_DARK = RGBColor(0x0F, 0x17, 0x2A)      # slate near-black
BRAND_MUTED = RGBColor(0x64, 0x74, 0x8B)     # slate-500
BRAND_LIGHT = RGBColor(0x94, 0xA3, 0xB8)     # slate-400

doc = Document()

# --- Default styles ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Heading colors aligned with the brand
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = BRAND_DARK if level == 1 else BRAND_DEEP
    hs.font.bold = True

# --- Helpers ---
def add_table(headers, rows, widths_in=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(10)
    for row in rows:
        r = t.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = str(val)
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if widths_in:
        for row in t.rows:
            for i, w in enumerate(widths_in):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def add_callout(text, kind="tip"):
    """Add a single-cell shaded callout box ('Tip' / 'Watch out' / 'Note')."""
    label_map = {
        "tip": ("Tip", "DBEAFE"),
        "warn": ("Watch out", "FEF3C7"),
        "note": ("Note", "F1F5F9"),
    }
    label, fill = label_map.get(kind, label_map["note"])
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    # Shading
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): fill,
    })
    tcPr.append(shd)
    cell.text = ""
    p = cell.paragraphs[0]
    label_run = p.add_run(f"{label}  ")
    label_run.bold = True
    label_run.font.color.rgb = BRAND_DEEP
    body_run = p.add_run(text)
    body_run.font.size = Pt(10.5)
    cell.width = Inches(6.5)
    doc.add_paragraph()


def add_steps(steps):
    """Numbered list of action steps using the built-in List Number style."""
    for s in steps:
        p = doc.add_paragraph(s, style="List Number")
        p.paragraph_format.space_after = Pt(4)


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)


def section_break():
    doc.add_page_break()


def hr_paragraph():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "1E40AF",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


# ========================================================================
# TITLE PAGE
# ========================================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ReportStack")
run.font.size = Pt(40)
run.bold = True
run.font.color.rgb = BRAND_DARK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("User Guide")
run.font.size = Pt(22)
run.font.color.rgb = BRAND_DEEP

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("Common workflows for QA teams, leads, stakeholders, and DevOps")
run.font.size = Pt(13)
run.font.color.rgb = BRAND_MUTED

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    f"Version 1.0\n{datetime.date.today().strftime('%B %d, %Y')}\n\n"
    "Companion to ReportStack_Documentation.docx (system reference)."
)
run.font.size = Pt(11)
run.font.color.rgb = BRAND_LIGHT

section_break()

# ========================================================================
# TABLE OF CONTENTS
# ========================================================================
doc.add_heading("Contents", level=1)
toc_items = [
    "1. How to use this guide",
    "2. Personas at a glance",
    "3. Scenario — QA engineer triages a failed run",
    "4. Scenario — QA lead sets up a new project",
    "5. Scenario — Stakeholder checks status via Telegram",
    "6. Scenario — Wiring a CI/CD pipeline",
    "7. Quick reference",
    "8. Glossary",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

section_break()

# ========================================================================
# 1. HOW TO USE THIS GUIDE
# ========================================================================
doc.add_heading("1. How to use this guide", level=1)
doc.add_paragraph(
    "This guide walks through the day-to-day workflows that make ReportStack useful "
    "for a QA team. Each scenario is a self-contained recipe — read only the ones "
    "that match your role, or skim the table of contents for what you need."
)
doc.add_paragraph(
    "If you are looking for the technical reference (architecture, data model, REST API, "
    "deployment), open ReportStack_Documentation.docx instead. This guide assumes the "
    "platform is already deployed and reachable at your team's ReportStack URL."
)

doc.add_heading("Conventions used here", level=2)
add_bullets([
    "URLs in monospace refer to routes inside the React frontend (for example, /launches/482).",
    "Code blocks show shell commands or pytest invocations you can copy directly.",
    "Side notes labelled Tip, Watch out, or Note give optional context — none of them are required to complete a scenario.",
    "Persona names (Maya, Jordan, Alex, Sam) are illustrative; substitute your own teammates.",
])

section_break()

# ========================================================================
# 2. PERSONAS
# ========================================================================
doc.add_heading("2. Personas at a glance", level=1)
doc.add_paragraph(
    "ReportStack is designed for four common roles. Each scenario in this guide "
    "is anchored to one of these personas, but the workflows generalise."
)

add_table(
    ["Persona", "Role", "Primary goal", "Lives in"],
    [
        ["Maya", "QA Engineer", "Triage failing tests and decide whether to file a defect", "Launch detail page; AI analysis panel"],
        ["Jordan", "QA Lead", "Set up new projects, manage the team, configure retention and notifications", "Settings, Members, Dashboards"],
        ["Alex", "Engineering stakeholder", "Stay informed about release readiness without opening the dashboard", "Telegram bot"],
        ["Sam", "DevOps engineer", "Wire test runs from CI/CD into ReportStack so reporting happens automatically", "Jenkins / GitHub Actions; pytest plugin"],
    ],
    widths_in=[1.0, 1.4, 2.7, 1.9],
)

section_break()

# ========================================================================
# 3. SCENARIO 1 — TRIAGE A FAILED RUN
# ========================================================================
doc.add_heading("3. Scenario — QA engineer triages a failed run", level=1)
doc.add_paragraph().add_run(
    "Maya, QA Engineer."
).italic = True

doc.add_heading("Goal", level=2)
doc.add_paragraph(
    "Maya saw a Slack notification that nightly regression failed. She wants to find "
    "out which tests broke, decide whether each failure is a real product bug or a flaky "
    "automation bug, and file the bugs to the engineering team."
)

doc.add_heading("Prerequisites", level=2)
add_bullets([
    "The launch already finished (or is finishing). ReportStack polls IN_PROGRESS launches every 5 seconds, so a near-real-time read is fine.",
    "Maya has at least the Member role on the project.",
    "The AI analyzer is enabled in Settings → Analyzer (it is by default).",
])

doc.add_heading("Steps", level=2)
add_steps([
    "Open the dashboard at /launches and click the failing launch (status badge shows Failed).",
    "On the launch detail page, the Tests tab opens by default. Sort or filter by Status = FAILED to focus on the broken tests.",
    "Click the first failing test name. The TestDetail page opens with a two-column layout: the error message + stack trace + log viewer on the left, and the AI analysis panel + defect selector on the right.",
    "Read the AI analysis panel (right rail). It classifies the failure as Product Bug, Automation Bug, System Issue, or To Investigate, and shows a confidence score plus a one-line reasoning.",
    "Decide whether to accept the AI's classification: if you agree, leave the Source = AI Auto badge in place; if not, click the defect selector dropdown, pick the right type (PB, AB, SI, ND, TI), and ReportStack records the override with Source = Manual.",
    "Add a Defect: in the right rail, click Add defect, write a one-line summary (\"login button missing in mobile viewport\"), set status to Open, optionally paste a Jira/Linear link, and save.",
    "Open Comments in the same page and leave a note for the team if context is needed (\"reproduces only with build 9241\").",
    "Repeat for each failing test. ReportStack remembers your overrides so the same flake will be auto-suggested as Automation Bug next time.",
])

add_callout(
    "Hover any log line to see the timestamp. Click a screenshot thumbnail to open it full-size in a modal — it loads from MinIO via /attachments/{id}/file.",
    "tip",
)

doc.add_heading("If you have many failures from the same root cause", level=2)
doc.add_paragraph(
    "Use the bulk-update endpoints from the launch detail toolbar (Bulk → Set defect type) or the bulk-analyze button to re-run AI classification across multiple selected failures. The retry-aware history strip on the test page also tells you whether this same test has been flaky in recent runs."
)

doc.add_heading("Cross-launch context", level=2)
doc.add_paragraph(
    "From any TestDetail page, the History strip at the bottom shows the same test's outcome across recent launches. Three reds in a row is a real bug; alternating green/red is a flake — that's usually enough to decide PB vs AB without reading another stack trace."
)

section_break()

# ========================================================================
# 4. SCENARIO 2 — SET UP A NEW PROJECT
# ========================================================================
doc.add_heading("4. Scenario — QA lead sets up a new project", level=1)
doc.add_paragraph().add_run(
    "Jordan, QA Lead."
).italic = True

doc.add_heading("Goal", level=2)
doc.add_paragraph(
    "Jordan inherited a new product line and wants ReportStack ready before the QA "
    "engineers point their pipelines at it. By the end of this scenario the project has "
    "retention rules, notification routes, defect-type colors, and a team roster."
)

doc.add_heading("Prerequisites", level=2)
add_bullets([
    "Admin role on the ReportStack instance. The first registered user is auto-promoted to ADMIN.",
    "An invite list of teammates' email addresses.",
    "Optional: a Slack webhook or a Telegram chat ID if you want broadcast notifications.",
])

doc.add_heading("Steps", level=2)
add_steps([
    "Sign in. The login surface uses the brand lockup; if you are first, register and you become Admin automatically.",
    "Open Settings (sidebar → Project → Settings). The first tab is General — set the project name, the inactivity timeout, and retention windows for launches, logs, and attachments. Pick numbers that match your audit and storage budget.",
    "Switch to the Notifications tab. Toggle the master switch on, configure email/Slack/Telegram channels, then create rules: pick a trigger (for example, Launch failed), recipients, and any attribute filters (tag = nightly).",
    "Switch to the Defect types tab. There are five default groups (PB, AB, SI, ND, TI). Edit colors and abbreviations to match how your team already talks about defects, or add subtypes (for example, AB → Selector flake) so the AI's auto-classifier has crisper buckets.",
    "Switch to the Analyzer tab and review the four sub-tabs (Index Settings, Auto-Analysis, Similar Items, Unique Errors). Adjust the auto-analysis match slider if you want the LLM to be more conservative.",
    "Open Members. Click Add member, enter a teammate's name, email, and role. Roles are Admin, Manager, Member, Viewer. Repeat for the team.",
    "Trigger a smoke launch from your test runner so the dashboard has data. The pytest plugin is the fastest path — see Scenario 6 for the exact CLI.",
    "Open Dashboards. The default dashboard is empty; click Add widget to drop in Launch Statistics Area, Pass Rate Pie, and Most Failed Test Cases. The widget grid is responsive and saves automatically.",
])

add_callout(
    "Retention rules are honoured by a daily cleanup daemon. Setting them too aggressively can lose forensic data — start permissive (90 days for launches, 30 days for attachments) and tighten later.",
    "warn",
)

section_break()

# ========================================================================
# 5. SCENARIO 3 — TELEGRAM BOT
# ========================================================================
doc.add_heading("5. Scenario — Stakeholder checks status via Telegram", level=1)
doc.add_paragraph().add_run(
    "Alex, Engineering Stakeholder."
).italic = True

doc.add_heading("Goal", level=2)
doc.add_paragraph(
    "Alex wants a one-line read on whether tonight's nightly run is healthy without "
    "opening a browser. ReportStack ships a Telegram bot powered by Claude that knows "
    "the project and can both run commands and answer free-text questions."
)

doc.add_heading("Prerequisites", level=2)
add_bullets([
    "The bot is deployed (docker compose up bot) and reachable on Telegram.",
    "Alex's Telegram user ID is in the ALLOWED_USER_IDS env var. If the list is empty the bot accepts everyone — fine for an internal team, risky on the open Telegram network.",
    "An ANTHROPIC_API_KEY is set so the agent can use Claude.",
])

doc.add_heading("What you can ask", level=2)
add_table(
    ["Command", "What it does"],
    [
        ["/status", "Backend health check + total launch count."],
        ["/launches", "Recent eight launches with pass rates."],
        ["/launch <id>", "Detailed launch info plus AI defect summary."],
        ["/failures <id>", "Failing tests with short error previews."],
        ["/analyze <id>", "Trigger AI failure classification on a launch."],
        ["/clear", "Reset the conversation history for fresh context."],
        ["Free text", "Routed to the Claude agent. Examples below."],
    ],
    widths_in=[1.6, 5.4],
)

doc.add_heading("Conversational examples", level=2)
doc.add_paragraph(
    "The bot has six tools (system health, launches list, launch detail, failed tests, "
    "analysis summary, trigger analysis). It chooses which to call based on the question."
)
add_bullets([
    "“Is the nightly green?” — bot calls get_launches, summarises the most recent nightly-tagged run.",
    "“Why did launch 482 fail?” — bot calls get_launch_detail and get_analysis_summary, lists the top defect types.",
    "“Re-run analysis on 482, the AI confidence looks low.” — bot calls trigger_analysis after confirming.",
    "“What's flaky this week?” — bot reads the most-failed endpoint and returns the top offenders.",
])

add_callout(
    "The bot keeps the last 20 turns of conversation history per user (MAX_HISTORY env var). Send /clear at the start of a new investigation if you want a clean slate.",
    "tip",
)

section_break()

# ========================================================================
# 6. SCENARIO 4 — CI/CD INTEGRATION
# ========================================================================
doc.add_heading("6. Scenario — Wiring a CI/CD pipeline", level=1)
doc.add_paragraph().add_run(
    "Sam, DevOps Engineer."
).italic = True

doc.add_heading("Goal", level=2)
doc.add_paragraph(
    "Sam wants every pipeline run to publish results to ReportStack automatically, "
    "trigger AI analysis on failure, and fail the build only when the launch finishes "
    "with a Failed or Stopped status."
)

doc.add_heading("Prerequisites", level=2)
add_bullets([
    "ReportStack reachable from the build agents on its API URL (default http://reports.local:8000/api/v1).",
    "An API token for the integration user (create one in ReportStack → /auth/users).",
    "Tests run with pytest. Other frameworks can use the REST API directly — see ReportStack_Documentation.docx, section 6.",
])

doc.add_heading("Step 1 — Install the pytest plugin", level=2)
doc.add_paragraph("On every build agent, install the plugin from the repo:")
doc.add_paragraph(
    "pip install -e plugins/pytest-automation-reports/"
).style = doc.styles.add_style("Code", 1) if "Code" not in [s.name for s in doc.styles] else doc.styles["Code"]

doc.add_heading("Step 2 — Run pytest with reporting flags", level=2)
add_bullets([
    "--ar-url — base URL for the ReportStack API (required).",
    "--ar-launch-name — human-readable launch name (\"Nightly regression\").",
    "--ar-launch-description — extra context (\"build 9241, branch main\").",
    "--ar-auto-analyze — kick off AI analysis as soon as the launch finishes.",
    "--ar-tag — tag the launch (use it to group nightly runs vs feature branches).",
])

doc.add_heading("Step 3 — Jenkins example", level=2)
doc.add_paragraph(
    "A minimal Jenkins declarative pipeline that publishes test results to ReportStack:"
)
add_table(
    ["Stage", "Shell"],
    [
        ["Install", "pip install -r requirements.txt && pip install -e plugins/pytest-automation-reports/"],
        ["Test",
         "pytest tests/ \\\n"
         "  --ar-url=$AR_URL \\\n"
         "  --ar-launch-name=\"$JOB_NAME #$BUILD_NUMBER\" \\\n"
         "  --ar-launch-description=\"build $GIT_COMMIT\" \\\n"
         "  --ar-tag=nightly --ar-auto-analyze"],
        ["Notify",
         "if ! pytest --exit-code; then\n"
         "  curl -X POST -H \"Authorization: Bearer $AR_TOKEN\" $AR_URL/launches/$LAUNCH_ID/analyze\n"
         "fi"],
    ],
    widths_in=[1.0, 5.5],
)

doc.add_heading("Step 4 — GitHub Actions example", level=2)
doc.add_paragraph(
    "Same idea expressed as a workflow step:"
)
add_table(
    ["Step", "YAML"],
    [
        ["Run tests",
         "- name: Run tests\n"
         "  run: |\n"
         "    pytest tests/ \\\n"
         "      --ar-url=${{ secrets.AR_URL }} \\\n"
         "      --ar-launch-name=\"${{ github.workflow }} #${{ github.run_number }}\" \\\n"
         "      --ar-tag=ci --ar-auto-analyze"],
    ],
    widths_in=[1.0, 5.5],
)

doc.add_heading("Storing the API URL and token", level=2)
add_bullets([
    "Use Jenkins Credentials (kind: Secret text) for AR_URL and AR_TOKEN; reference them as withCredentials in the pipeline.",
    "On GitHub Actions use repository secrets (Settings → Secrets and variables → Actions). Never hard-code tokens in YAML.",
    "Rotate the integration token quarterly. ReportStack does not rotate it for you.",
])

add_callout(
    "Set --ar-launch-name to a stable, predictable string (job + build number) so the dashboard's launch list groups runs sensibly. Free-form names with timestamps make the table noisy.",
    "tip",
)

section_break()

# ========================================================================
# 7. QUICK REFERENCE
# ========================================================================
doc.add_heading("7. Quick reference", level=1)

doc.add_heading("Defect type abbreviations", level=2)
add_table(
    ["Abbr", "Type", "Meaning"],
    [
        ["PB", "Product bug", "A real bug in the application under test."],
        ["AB", "Automation bug", "Selector drift, timing, or test-code flake."],
        ["SI", "System issue", "Infra failure (network, container, dependency)."],
        ["ND", "No defect", "Test was correct; the failure was a false alarm."],
        ["TI", "To investigate", "Default when AI confidence is low (< 0.4)."],
    ],
    widths_in=[0.7, 1.7, 4.1],
)

doc.add_heading("Useful frontend routes", level=2)
add_table(
    ["Route", "Purpose"],
    [
        ["/", "Dashboard overview metrics."],
        ["/launches", "Paginated launch list with filters."],
        ["/launches/:id", "Launch detail (Tests + Comments tabs)."],
        ["/launches/:id/items/:itemId", "Test detail (error, logs, AI panel, history)."],
        ["/dashboards", "Custom widget grids."],
        ["/trends", "Cross-launch trend analysis."],
        ["/members", "Team roster."],
        ["/settings", "9-tab project settings."],
    ],
    widths_in=[2.4, 4.1],
)

doc.add_heading("REST endpoints worth knowing", level=2)
add_table(
    ["Method + path", "Purpose"],
    [
        ["POST /launches/", "Create a launch."],
        ["PUT /launches/{id}/finish", "Finish a launch with final status."],
        ["POST /launches/{id}/items/batch", "Report many test results at once."],
        ["POST /launches/{id}/analyze", "Trigger AI analysis on the launch."],
        ["POST /launches/{id}/items/{item_id}/attachments", "Upload a screenshot (multipart)."],
        ["GET /items/most-failed", "Cross-launch flakiness aggregation."],
    ],
    widths_in=[3.0, 3.5],
)

section_break()

# ========================================================================
# 8. GLOSSARY
# ========================================================================
doc.add_heading("8. Glossary", level=1)
add_table(
    ["Term", "Meaning"],
    [
        ["Launch", "A single test run, typically one CI build's worth of tests."],
        ["Test item", "One test case inside a launch; status PASSED/FAILED/SKIPPED/ERROR."],
        ["Defect", "A linked issue (with type and status) attached to a failing test item."],
        ["Defect type", "Classification bucket — PB, AB, SI, ND, TI — see Quick reference."],
        ["AI analysis", "Local LLM (Ollama) classification of a failed test, with confidence score and reasoning."],
        ["Tag", "Free-form label on a launch (\"nightly\", \"feature/x\"), used for filtering."],
        ["Retention", "Time-based cleanup window for launches, logs, and attachments."],
        ["History strip", "Cross-launch outcome line on the test detail page; reveals flakes."],
        ["Widget", "A single chart or table on a Dashboard page."],
    ],
    widths_in=[1.7, 4.8],
)

hr_paragraph()
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = final.add_run(
    f"Generated {datetime.date.today().strftime('%B %d, %Y')}  ·  "
    "Companion to ReportStack_Documentation.docx"
)
run.font.size = Pt(9)
run.font.color.rgb = BRAND_LIGHT

# ========================================================================
# SAVE
# ========================================================================
output_path = "/Users/rashadjafarzadeh/Projects/automation-reports/ReportStack_User_Guide.docx"
doc.save(output_path)
print(f"User guide saved to: {output_path}")
