"""Generate ReportStack Linux integration prerequisites as a Word document.

Mirrors deploy/PREREQUISITES.md so the same checklist is portable across
devices that don't have the repo checked out.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

BRAND_DEEP = RGBColor(0x1E, 0x40, 0xAF)
BRAND_DARK = RGBColor(0x0F, 0x17, 0x2A)
BRAND_MUTED = RGBColor(0x64, 0x74, 0x8B)
BRAND_LIGHT = RGBColor(0x94, 0xA3, 0xB8)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = BRAND_DARK if level == 1 else BRAND_DEEP
    hs.font.bold = True


def add_table(headers, rows, widths_in=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(10)
    for row in rows:
        r = t.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = str(val)
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if widths_in:
        for row in t.rows:
            for i, w in enumerate(widths_in):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def add_callout(text, kind="tip"):
    label_map = {
        "tip": ("Tip", "DBEAFE"),
        "warn": ("Watch out", "FEF3C7"),
        "note": ("Note", "F1F5F9"),
    }
    label, fill = label_map.get(kind, label_map["note"])
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): fill,
    })
    tcPr.append(shd)
    cell.text = ""
    p = cell.paragraphs[0]
    label_run = p.add_run(f"{label}  ")
    label_run.bold = True
    label_run.font.color.rgb = BRAND_DEEP
    body_run = p.add_run(text)
    body_run.font.size = Pt(10.5)
    cell.width = Inches(6.5)
    doc.add_paragraph()


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)


def add_code_block(text):
    """Render a code block as a single shaded cell with monospace text."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F1F5F9",
    })
    tcPr.append(shd)
    cell.text = ""
    for i, line in enumerate(text.splitlines()):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_DARK
    cell.width = Inches(6.5)
    doc.add_paragraph()


def section_break():
    doc.add_page_break()


# ============================================================
# COVER
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ReportStack")
run.font.size = Pt(40)
run.bold = True
run.font.color.rgb = BRAND_DARK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Linux integration prerequisites")
run.font.size = Pt(22)
run.font.color.rgb = BRAND_DEEP

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("Tools and apps checklist for deploying to twd00030 (Ubuntu 24.04 LTS)")
run.font.size = Pt(13)
run.font.color.rgb = BRAND_MUTED

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    f"Version 1.0\n{datetime.date.today().strftime('%B %d, %Y')}\n\n"
    "Companion to ReportStack_Documentation.docx and ReportStack_User_Guide.docx."
)
run.font.size = Pt(11)
run.font.color.rgb = BRAND_LIGHT

section_break()

# ============================================================
# TOC
# ============================================================
doc.add_heading("Contents", level=1)
toc = [
    "1. How to read this document",
    "2. Host system tools",
    "3. Docker images",
    "4. NVIDIA GPU (optional)",
    "5. End-to-end test suite",
    "6. pytest plugin",
    "7. Telegram bot env vars",
    "8. Secrets and config",
    "9. Network ports",
    "10. Jenkins integration (optional)",
    "11. Quick deploy — the happy path",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
section_break()

# ============================================================
# 1. INTRO
# ============================================================
doc.add_heading("1. How to read this document", level=1)
doc.add_paragraph(
    "This is the inventory of every tool, library, image, environment variable, "
    "and secret that has to exist on twd00030 before ReportStack runs. The host "
    "is Ubuntu 24.04.4 LTS. Items already present per the cicd skill are flagged "
    "with a check; items needing fresh install are flagged with a download arrow."
)
doc.add_paragraph(
    "If you have shell access on a different Linux machine (laptop, CI agent, a "
    "QA workstation), the same list applies — substitute the apt commands for "
    "your distro's package manager where relevant."
)
add_callout(
    "Status legend used in tables: ✓ means already installed on twd00030, "
    "⇩ means install needed, and ⚠ means optional capability.",
    "note",
)

# ============================================================
# 2. HOST TOOLS
# ============================================================
doc.add_heading("2. Host system tools", level=1)
doc.add_paragraph(
    "Base toolchain. Everything else builds on top of these. The host already "
    "has all of them — verify with the install commands below if you're "
    "deploying to a different machine."
)
add_table(
    ["Tool", "Why we need it", "Status"],
    [
        ["Docker Engine ≥ 24", "Run every service except optionally the test runner", "✓ 29.4.1"],
        ["Docker Compose v2", "docker compose ... (space, not hyphen)", "✓"],
        ["git", "Pull and push the repo", "✓ 2.43"],
        ["Python 3.12", "Matches backend; needed for seed_data.py and ad-hoc uvicorn runs", "✓ 3.12.3"],
        ["Node.js 18+ + npm", "Frontend build, Playwright, pytest plugin install", "✓ 25.5 / 11.8"],
        ["curl, wget, jq, unzip", "Smoke checks, debugging, asset downloads", "✓"],
        ["gcc, make", "Native deps when pip wheels miss (psycopg2 fallback, lxml)", "✓"],
    ],
    widths_in=[1.6, 3.2, 1.7],
)
add_callout(
    "Don't install Java on the host. Jenkins ships its own JDK in its container — "
    "host-installed Java will conflict.",
    "warn",
)

# ============================================================
# 3. DOCKER IMAGES
# ============================================================
doc.add_heading("3. Docker images", level=1)
doc.add_paragraph(
    "These are pulled or built by docker compose up -d --build. The first run "
    "downloads a few GB; subsequent runs are diff-only."
)
add_table(
    ["Image", "Source", "Approx size"],
    [
        ["postgres:16-alpine", "Docker Hub", "~250 MB"],
        ["minio/minio:latest", "Docker Hub", "~200 MB"],
        ["ollama/ollama:latest", "Docker Hub", "~700 MB"],
        ["automation-reports_backend", "built from backend/Dockerfile", "~250 MB"],
        ["automation-reports_frontend", "built from frontend/Dockerfile (Node → nginx)", "~50 MB"],
        ["automation-reports_bot", "built from bot/Dockerfile", "~180 MB"],
    ],
    widths_in=[2.4, 2.6, 1.5],
)
doc.add_paragraph("After the stack is up, pull the LLM weights once:")
add_code_block("docker compose exec ollama ollama pull mistral:7b")
doc.add_paragraph(
    "That adds about 4 GB on disk. Skip it only if you accept the analyzer's "
    "TO_INVESTIGATE fallback for every failure (the May 2026 fix in "
    "ai_analyzer.py makes this safe — no crashes, just less useful "
    "classifications)."
)

# ============================================================
# 4. NVIDIA GPU
# ============================================================
doc.add_heading("4. NVIDIA GPU (optional)", level=1)
doc.add_paragraph(
    "Only relevant if the box has a GPU and you want fast Ollama inference. "
    "Skip this section for CPU-only deploys."
)
add_table(
    ["Tool", "Install command"],
    [
        ["NVIDIA driver", "sudo apt install nvidia-driver-550"],
        ["nvidia-container-toolkit", "Follow the official install guide; configure with sudo nvidia-ctk runtime configure --runtime=docker"],
        ["Restart Docker", "sudo systemctl restart docker"],
    ],
    widths_in=[2.0, 4.5],
)
doc.add_paragraph(
    "The ollama service block in docker-compose.yml already declares "
    "deploy.resources.reservations.devices for the GPU. Remove that block "
    "for CPU-only runs so compose doesn't reject the file."
)

# ============================================================
# 5. E2E SUITE
# ============================================================
doc.add_heading("5. End-to-end test suite", level=1)
doc.add_paragraph(
    "Playwright drives a Chromium browser against the live frontend. Required "
    "on any machine where the suite is run — typically a CI agent or a dev box."
)
doc.add_heading("Install steps", level=2)
add_code_block(
    "cd tests/e2e\n"
    "npm install\n"
    "npx playwright install chromium\n"
    "npx playwright install-deps"
)
doc.add_heading("Explicit Ubuntu 24.04 system libs", level=2)
doc.add_paragraph(
    "If install-deps complains, install these directly. The t64 suffix is "
    "Ubuntu 24.04's time64 transition rename — older docs reference the names "
    "without it."
)
add_code_block(
    "sudo apt install -y \\\n"
    "  libnss3 libatk1.0-0t64 libatk-bridge2.0-0t64 libcups2t64 \\\n"
    "  libxcomposite1 libxdamage1 libxfixes3 libxrandr2 libgbm1 \\\n"
    "  libxkbcommon0 libpango-1.0-0 libcairo2 libasound2t64"
)
doc.add_heading("Headed mode in CI (optional)", level=2)
doc.add_paragraph(
    "If you ever need a real X display, install Xvfb and wrap the test run:"
)
add_code_block("sudo apt install -y xvfb\nxvfb-run npm test")

# ============================================================
# 6. PYTEST PLUGIN
# ============================================================
doc.add_heading("6. pytest plugin", level=1)
doc.add_paragraph(
    "Required on any CI worker or dev box that runs real test suites and reports "
    "results to ReportStack."
)
add_code_block("pip install -e plugins/pytest-automation-reports/")
doc.add_paragraph(
    "Depends on pytest ≥ 7.0, requests, and Pillow (screenshots). Selenium "
    "and Playwright Python bindings are picked up at runtime if your tests use "
    "them — the plugin doesn't pin them."
)

# ============================================================
# 7. TELEGRAM BOT ENV
# ============================================================
doc.add_heading("7. Telegram bot env vars", level=1)
doc.add_paragraph(
    "The bot/ container is in the compose file but won't start cleanly without "
    "these. Drop them in .env next to docker-compose.yml."
)
add_table(
    ["Env var", "Where to get it"],
    [
        ["TELEGRAM_BOT_TOKEN", "@BotFather on Telegram"],
        ["ANTHROPIC_API_KEY", "console.anthropic.com"],
        ["CLAUDE_MODEL", "Optional, default claude-sonnet-4-20250514"],
        ["ALLOWED_USER_IDS", "Comma-separated Telegram numeric IDs; empty = open to all"],
        ["AR_BACKEND_URL", "http://backend:8000/api/v1 (Docker DNS)"],
        ["AR_FRONTEND_URL", "External dashboard URL (used in clickable bot links)"],
    ],
    widths_in=[2.2, 4.3],
)

# ============================================================
# 8. SECRETS / CONFIG
# ============================================================
doc.add_heading("8. Secrets and config", level=1)
add_table(
    ["File / var", "Purpose"],
    [
        [".env (root)", "Compose pulls this automatically; never commit it"],
        ["JWT_SECRET", "Signs auth tokens. Generate with openssl rand -hex 32."],
        ["MINIO_ACCESS_KEY / SECRET_KEY", "Default minioadmin works internal-only; rotate for anything internet-facing"],
        ["OLLAMA_BASE_URL", "http://ollama:11434 inside the Docker network"],
        ["DATABASE_URL", "postgresql://postgres:postgres@db:5432/automation_reports"],
    ],
    widths_in=[2.6, 3.9],
)
doc.add_paragraph("There is already an .env.example in the repo; copy from it.")

# ============================================================
# 9. PORTS
# ============================================================
doc.add_heading("9. Network ports", level=1)
add_table(
    ["Port", "Service", "Bind recommendation"],
    [
        ["3000", "Frontend (nginx)", "Public if behind internal proxy"],
        ["8000", "Backend (FastAPI)", "Public if same-origin from frontend"],
        ["5432", "PostgreSQL", "127.0.0.1 only"],
        ["9000 / 9001", "MinIO (S3 API + console)", "127.0.0.1 or VPN-only"],
        ["11434", "Ollama", "127.0.0.1 only"],
    ],
    widths_in=[1.0, 2.5, 3.0],
)
doc.add_paragraph(
    "If the company internal proxy fronts the dashboard, only 3000 needs to be "
    "externally reachable; everything else stays on the loopback or Docker "
    "bridge network."
)

# ============================================================
# 10. JENKINS
# ============================================================
doc.add_heading("10. Jenkins integration (optional)", level=1)
doc.add_paragraph(
    "Already covered in cicd-skill.md — Jenkins runs in its own Docker container "
    "with its own JDK. The CI prerequisites that touch ReportStack:"
)
add_bullets([
    "A Jenkins Secret-text credential AR_TOKEN, created via POST /api/v1/auth/users with an integration user.",
    "A Jenkins Secret-text credential AR_URL set to the internal API URL.",
    "A pipeline shell step that installs the pytest plugin and passes --ar-url=$AR_URL --ar-launch-name=\"$JOB_NAME #$BUILD_NUMBER\" --ar-auto-analyze.",
])
doc.add_paragraph(
    "The User Guide § Scenario 6 has a copy-pasteable Jenkinsfile fragment."
)

# ============================================================
# 11. QUICK DEPLOY
# ============================================================
doc.add_heading("11. Quick deploy — the happy path", level=1)
doc.add_paragraph(
    "Once the prereqs are satisfied, the sequence below brings the full stack up."
)
add_code_block(
    "ssh user@twd00030\n"
    "cd ~/automation-reports\n"
    "git pull origin main\n"
    "\n"
    "# secrets — edit values\n"
    "cp .env.example .env\n"
    "${EDITOR:-vim} .env\n"
    "\n"
    "# build images and bring everything up\n"
    "docker compose up -d --build\n"
    "docker compose exec ollama ollama pull mistral:7b   # one-off\n"
    "\n"
    "# run migrations\n"
    "docker compose exec backend alembic upgrade head\n"
    "\n"
    "# health check\n"
    "curl -fsS http://localhost:8000/api/v1/launches/\n"
    "curl -fsS http://localhost:3000/\n"
    "\n"
    "# (optional) run the e2e suite against the live stack\n"
    "cd tests/e2e\n"
    "npm install && npx playwright install chromium && npx playwright install-deps\n"
    "RS_FRONTEND_URL=http://localhost:3000 \\\n"
    "  RS_API_URL=http://localhost:8000/api/v1 \\\n"
    "  npm test"
)
add_callout(
    "If any of the curls fail, see cicd-skill.md § Common operations / Gotchas, "
    "or run docker compose logs --tail 100 <service-name> to inspect.",
    "tip",
)

# ============================================================
# SAVE
# ============================================================
output_path = "/Users/rashadjafarzadeh/Projects/automation-reports/ReportStack_Linux_Prerequisites.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
